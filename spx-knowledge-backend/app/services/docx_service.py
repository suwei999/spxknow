import os
from typing import Any, Dict, List, Optional, Iterable, Tuple
import zipfile
import tempfile
import shutil
import re

from sqlalchemy.orm import Session

from app.core.logging import logger


NAMESPACES = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}


def _clear_element(paragraphs: List[Any]) -> None:
    if not paragraphs:
        return
    for para in list(paragraphs):
        if para.text and para.text.strip():
            para.text = ''
        if para.runs:
            for run in para.runs:
                run.text = ''


def _clear_tables(tables: List[Any]) -> None:
    if not tables:
        return
    for table in list(tables):
        for row in table.rows:
            for cell in row.cells:
                cell.text = ''


def _is_toc_line(text: str) -> bool:
    if not text:
        return False
    stripped = text.strip()
    if not stripped:
        return True
    lower = stripped.lower()
    if lower.startswith('目录') or lower.startswith('table of contents'):
        return True
    # 匹配形如“章节名 ...... 12”的目录项，兼容中文省略号与圆点
    if re.search(r"[\.·\u2026]{2,}\s*\d+$", stripped):
        return True
    return False


def _is_toc_paragraph(paragraph: Any) -> bool:
    """综合段落样式、内容与 XML 指令判断是否为目录条目"""
    try:
        style_name = getattr(getattr(paragraph, 'style', None), 'name', '') or ''
        if style_name and style_name.upper().startswith('TOC'):
            return True
    except Exception:
        pass

    try:
        runs_text = ''.join((run.text or '') for run in getattr(paragraph, 'runs', []) if getattr(run, 'text', None))
    except Exception:
        runs_text = ''

    text = (paragraph.text or '').strip() if hasattr(paragraph, 'text') else ''
    combined_text = (runs_text or text).strip()
    if _is_toc_line(combined_text):
        return True

    try:
        xml_repr = paragraph._p.xml  # type: ignore[attr-defined]
        if xml_repr and 'TOC' in xml_repr.upper():
            return True
    except Exception:
        pass

    return False


def _flatten_table_text(table: Any) -> str:
    lines: List[str] = []
    try:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = (cell.text or '').strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                lines.append(' '.join(cells))
    except Exception:
        return ''
    return ' '.join(lines).strip()


def _is_toc_table(table: Any) -> bool:
    text = _flatten_table_text(table)
    if not text:
        return False
    normalized = text.replace(' ', '').lower()
    if normalized.startswith('目录') or normalized.startswith('tableofcontents'):
        return True
    if re.search(r"[\.·\u2026]{2,}\s*\d+", text):
        dot_matches = re.findall(r"[\.·\u2026]{2,}\s*\d+", text)
        if len(dot_matches) >= 2:
            return True
    return False


def _remove_toc_content_controls(doc: Any) -> int:
    removed = 0
    try:
        nsmap = doc._element.nsmap  # type: ignore[attr-defined]
        toc_nodes = list(doc._element.xpath('.//w:sdt', namespaces=nsmap))  # type: ignore[attr-defined]
        for sdt in toc_nodes:
            try:
                text_nodes = sdt.xpath('.//w:t/text()', namespaces=nsmap)  # type: ignore[attr-defined]
                combined = ''.join(text_nodes).strip()
                normalized = combined.replace(' ', '').lower()
                if normalized.startswith('目录') or normalized.startswith('tableofcontents') or 'toc' in normalized:
                    parent = sdt.getparent()
                    if parent is not None:
                        parent.remove(sdt)
                        removed += 1
            except Exception:
                continue
    except Exception:
        return removed
    return removed


class DocxService:
    """
    DOCX 文档解析服务（完全本地，基于 python-docx），不依赖 Unstructured。
    输出结构与现有任务流兼容：
      - text_content: str
      - tables: [{element_index, table_data{cells/html/rows/columns}, table_text, page_number=None}]
      - images: [{data/bytes, element_index(None), page_number=None}]
      - text_element_index_map: [{element_index, element_type}]
      - filtered_elements_light: [{category, text, element_index}]
    """

    def __init__(self, db: Session):
        self.db = db

    @staticmethod
    def has_embedded_images(file_path: str) -> bool:
        try:
            if not os.path.exists(file_path):
                return False
            with zipfile.ZipFile(file_path, 'r') as zf:
                for name in zf.namelist():
                    if name.startswith('word/media/'):
                        return True
        except Exception as exc:
            logger.debug(f"[DOCX] 检查图片失败: {exc}")
        return False

    @staticmethod
    def sanitize_docx(file_path: str) -> str:
        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            logger.warning(f"[DOCX] sanitize_docx 导入 python-docx 失败: {exc}")
            return file_path

        try:
            doc = Document(file_path)
        except Exception as exc:
            logger.warning(f"[DOCX] sanitize_docx 打开失败: {exc}")
            return file_path

        try:
            for section in doc.sections:
                _clear_element(section.header.paragraphs)
                _clear_element(section.footer.paragraphs)
                if hasattr(section.header, 'tables'):
                    _clear_tables(section.header.tables)
                if hasattr(section.footer, 'tables'):
                    _clear_tables(section.footer.tables)
        except Exception as exc:
            logger.debug(f"[DOCX] sanitize_docx 清理页眉页脚失败: {exc}")

        try:
            processed = 0
            max_scan = min(len(doc.paragraphs), 200)
            idx = 0
            while idx < max_scan and idx < len(doc.paragraphs):
                para = doc.paragraphs[idx]
                if _is_toc_paragraph(para):
                    doc._element.body.remove(para._element)
                    processed += 1
                    max_scan = min(len(doc.paragraphs), 200)
                    continue
                idx += 1
            if processed:
                logger.debug(f"[DOCX] sanitize_docx 移除目录段落 {processed} 条")
        except Exception as exc:
            logger.debug(f"[DOCX] sanitize_docx 清理目录失败: {exc}")

        # 清理以表格呈现的目录
        try:
            removed_tables = 0
            for tbl in list(doc.tables):
                if _is_toc_table(tbl):
                    parent = tbl._element.getparent()
                    if parent is not None:
                        parent.remove(tbl._element)
                        removed_tables += 1
            if removed_tables:
                logger.debug(f"[DOCX] sanitize_docx 移除目录表格 {removed_tables} 个")
        except Exception as exc:
            logger.debug(f"[DOCX] sanitize_docx 清理目录表格失败: {exc}")

        # 清理内容控件形式的目录（Word 自动目录常见形式）
        try:
            removed_controls = _remove_toc_content_controls(doc)
            if removed_controls:
                logger.debug(f"[DOCX] sanitize_docx 移除目录内容控件 {removed_controls} 个")
        except Exception as exc:
            logger.debug(f"[DOCX] sanitize_docx 清理目录内容控件失败: {exc}")

        try:
            tmp_dir = tempfile.mkdtemp(prefix='docx_sanitized_')
            sanitized_path = os.path.join(tmp_dir, os.path.basename(file_path))
            doc.save(sanitized_path)
            return sanitized_path
        except Exception as exc:
            logger.warning(f"[DOCX] sanitize_docx 保存失败: {exc}")
            try:
                shutil.rmtree(tmp_dir)
            except Exception:
                pass
            return file_path

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            raise RuntimeError("缺少依赖 python-docx，请先安装: pip install python-docx") from e

        logger.info(f"[DOCX] 开始本地解析: {file_path}")
        doc = Document(file_path)

        ordered_elements: List[Dict[str, Any]] = []
        text_element_index_map: List[Dict[str, Any]] = []
        filtered_elements_light: List[Dict[str, Any]] = []
        text_content_parts: List[str] = []
        images_payload: List[Dict[str, Any]] = []
        tables: List[Dict[str, Any]] = []
        images_binary: List[Dict[str, Any]] = []

        element_index = 0
        doc_order = 0

        def _join_text_segments(segments: Iterable[str]) -> str:
            return ' '.join(seg.strip() for seg in segments if seg and seg.strip())

        def _append_text_segment(text: str, style_name: str) -> None:
            nonlocal element_index, doc_order
            normalized = (text or '').strip()
            if not normalized:
                return
            lower_style = (style_name or '').lower()
            is_title = False
            if style_name:
                if style_name.startswith('Heading') or 'title' in lower_style or '标题' in lower_style:
                    is_title = True
            category = 'Title' if is_title else 'NarrativeText'
            element_index += 1
            ordered_elements.append({
                'type': 'text',
                'text': normalized,
                'style': style_name or '',
                'category': category,
                'element_index': element_index,
                'doc_order': doc_order,
                'length': len(normalized),
            })
            text_content_parts.append(normalized)
            filtered_elements_light.append({
                'category': category,
                'text': normalized,
                'element_index': element_index,
                'doc_order': doc_order,
            })
            text_element_index_map.append({
                'element_index': element_index,
                'element_type': category,
                'doc_order': doc_order,
                'page_number': None,
                'coordinates': None,
            })
            doc_order += 1

        def _append_image_segment(r_id: str) -> None:
            nonlocal element_index, doc_order
            if not r_id:
                return
            part = doc.part.related_parts.get(r_id)
            if not part:
                return
            data = part.blob
            if not data:
                return
            element_index += 1
            image_ext = os.path.splitext(part.partname)[1] or ''
            image_meta = {
                'data': data,
                'bytes': data,
                'element_index': element_index,
                'doc_order': doc_order,
                'page_number': None,
                'coordinates': None,
                'rId': r_id,
                'image_ext': image_ext,
            }
            images_payload.append(image_meta)
            images_binary.append({
                'binary': data,
                'element_index': element_index,
                'page_number': None,
                'doc_order': doc_order,
            })
            ordered_elements.append({
                'type': 'image',
                'element_index': element_index,
                'doc_order': doc_order,
                'rId': r_id,
                'image_ext': image_ext,
            })
            doc_order += 1

        def _iter_paragraph_segments(paragraph) -> Iterable[Tuple[str, str]]:
            try:
                run_elements = paragraph._p.xpath('./w:r | ./w:hyperlink | ./w:fldSimple', namespaces=NAMESPACES)
            except Exception:
                try:
                    run_elements = paragraph._p.xpath('./w:r | ./w:hyperlink | ./w:fldSimple')
                except Exception:
                    run_elements = list(paragraph._p.iterchildren())

            for run in run_elements:
                targets = []
                if run.tag == qn('w:hyperlink'):
                    try:
                        targets = run.xpath('.//w:r', namespaces=NAMESPACES)
                    except Exception:
                        try:
                            targets = run.xpath('.//w:r')
                        except Exception:
                            targets = list(run.iterchildren())
                elif run.tag == qn('w:fldSimple'):
                    try:
                        targets = run.xpath('.//w:r', namespaces=NAMESPACES)
                    except Exception:
                        try:
                            targets = run.xpath('.//w:r')
                        except Exception:
                            targets = list(run.iterchildren())
                else:
                    targets = [run]

                for r in targets:
                    try:
                        text_nodes = r.xpath('.//w:t', namespaces=NAMESPACES)
                    except Exception:
                        try:
                            text_nodes = r.xpath('.//w:t')
                        except Exception:
                            text_nodes = []
                    texts = [t.text for t in text_nodes if t.text]
                    if texts:
                        yield ('text', ''.join(texts))
                    try:
                        blips = r.xpath('.//a:blip', namespaces=NAMESPACES)
                    except Exception:
                        try:
                            blips = r.xpath('.//a:blip')
                        except Exception:
                            blips = []
                    for blip in blips:
                        r_id = blip.get(qn('r:embed'))
                        if r_id:
                            yield ('image', r_id)

        try:
            from docx.oxml.ns import qn  # type: ignore
            from docx.text.paragraph import Paragraph  # type: ignore
            from docx.table import Table as DxTable  # type: ignore

            def iter_block_items(document):
                body = document.element.body
                for child in body.iterchildren():
                    tag = child.tag
                    if tag == qn('w:p'):
                        yield 'paragraph', Paragraph(child, document)
                    elif tag == qn('w:tbl'):
                        yield 'table', DxTable(child, document)

            for kind, item in iter_block_items(doc):
                if kind == 'paragraph':
                    style_name = getattr(getattr(item, 'style', None), 'name', '') or ''
                    text_segments: List[str] = []
                    for seg_type, payload in _iter_paragraph_segments(item):
                        if seg_type == 'text':
                            text_segments.append(payload)
                        elif seg_type == 'image':
                            if text_segments:
                                _append_text_segment(_join_text_segments(text_segments), style_name)
                                text_segments = []
                            _append_image_segment(payload)
                    if text_segments:
                        _append_text_segment(_join_text_segments(text_segments), style_name)
                elif kind == 'table':
                    cells: List[List[str]] = []
                    for row in item.rows:
                        row_values: List[str] = []
                        for cell in row.cells:
                            cell_text_segments: List[str] = []
                            for paragraph in cell.paragraphs:
                                cell_text_segments.append(paragraph.text or '')
                            cell_text = _join_text_segments(cell_text_segments)
                            row_values.append(cell_text)
                        if any((val or '').strip() for val in row_values):
                            cells.append(row_values)
                    if not cells:
                        continue
                    table_data = {
                        'cells': cells,
                        'rows': len(cells),
                        'columns': len(cells[0]) if cells and cells[0] else 0,
                        'structure': 'docx_extracted',
                        'html': None,
                    }
                    table_text = '\n'.join('\t'.join(str(cell or '') for cell in row) for row in cells)
                    element_index += 1
                    tables.append({
                        'element_index': element_index,
                        'table_data': table_data,
                        'table_text': table_text,
                        'page_number': None,
                        'doc_order': doc_order,
                    })
                    ordered_elements.append({
                        'type': 'table',
                        'element_index': element_index,
                        'doc_order': doc_order,
                        'table_data': table_data,
                        'table_text': table_text,
                        'page_number': None,
                    })
                    filtered_elements_light.append({
                        'category': 'Table',
                        'text': table_text,
                        'element_index': element_index,
                        'doc_order': doc_order,
                    })
                    doc_order += 1
        except Exception as exc:
            logger.warning(f"按顺序遍历段落/表格失败，降级为基础解析: {exc}")
            for paragraph in doc.paragraphs:
                text_value = (paragraph.text or '').strip()
                if not text_value:
                    continue
                style_name = getattr(getattr(paragraph, 'style', None), 'name', '') or ''
                _append_text_segment(text_value, style_name)
            for table in getattr(doc, 'tables', []):
                cells: List[List[str]] = []
                for row in table.rows:
                    row_values = [(_join_text_segments([cell.text])) for cell in row.cells]
                    if any((val or '').strip() for val in row_values):
                        cells.append(row_values)
                if not cells:
                    continue
                table_data = {
                    'cells': cells,
                    'rows': len(cells),
                    'columns': len(cells[0]) if cells and cells[0] else 0,
                    'structure': 'docx_extracted',
                    'html': None,
                }
                table_text = '\n'.join('\t'.join(str(cell or '') for cell in row) for row in cells)
                element_index += 1
                tables.append({
                    'element_index': element_index,
                    'table_data': table_data,
                    'table_text': table_text,
                    'page_number': None,
                    'doc_order': doc_order,
                })
                ordered_elements.append({
                    'type': 'table',
                    'element_index': element_index,
                    'doc_order': doc_order,
                    'table_data': table_data,
                    'table_text': table_text,
                    'page_number': None,
                })
                filtered_elements_light.append({
                    'category': 'Table',
                    'text': table_text,
                    'element_index': element_index,
                    'doc_order': doc_order,
                })
                doc_order += 1

        metadata = {
            'element_count': len(ordered_elements),
            'images_count': len(images_payload),
            'filter_stats': {},
        }

        parse_result: Dict[str, Any] = {
            'text_content': ('\n'.join(text_content_parts)).strip(),
            'tables': tables,
            'images': images_payload,
            'images_binary': images_binary,
            'ordered_elements': ordered_elements,
            'text_element_index_map': text_element_index_map,
            'filtered_elements_light': filtered_elements_light,
            'metadata': metadata,
            'is_converted_pdf': False,
            'converted_pdf_path': None,
        }
        logger.info(
            f"[DOCX] 解析完成: 元素={len(ordered_elements)}, 文本={sum(1 for e in ordered_elements if e['type']=='text')}, "
            f"表格={sum(1 for e in ordered_elements if e['type']=='table')}, 图片={len(images_payload)}"
        )
        return parse_result

    # ============== 分块（复用结构式分块逻辑） ==============
    def chunk_text(
        self,
        text: str,
        text_element_index_map: Optional[Dict[int, Dict[str, Any]]] = None,
        elements: Optional[List[Any]] = None
    ) -> List[Dict[str, Any]]:
        if not text or not text.strip():
            return []
        # 若提供 elements，按结构分块
        if elements:
            return self._chunk_by_structure(elements, text_element_index_map)
        # 否则简单段落分块
        chunks: List[Dict[str, Any]] = []
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        for i, para in enumerate(paragraphs):
            chunks.append({
                'content': para,
                'chunk_index': i,
                'element_index_start': None,
                'element_index_end': None,
                'section_id': None,
                'is_parent': False,
            })
        return chunks

    def _chunk_by_structure(
        self,
        elements: List[Any],
        text_element_index_map: Optional[Dict[int, Dict[str, Any]]] = None
    ) -> List[Dict[str, Any]]:
        # 严格控制单块大小：与 settings 对齐（默认 1000 / 1024）
        try:
            from app.config.settings import settings as _settings
            chunk_max = min(int(getattr(_settings, 'CHUNK_SIZE', 1000)), int(getattr(_settings, 'TEXT_EMBED_MAX_CHARS', 1024)))
            # 最小分块大小：避免产生太多小块（默认 500 字符，约为 chunk_max 的一半）
            chunk_min = int(getattr(_settings, 'CHUNK_MIN_SIZE', chunk_max // 2))
        except Exception:
            chunk_max = 1000
            chunk_min = 500
        chunks: List[Dict[str, Any]] = []
        current_chunk_texts: List[str] = []
        current_chunk_start: Optional[int] = None
        current_chunk_end: Optional[int] = None
        current_length = 0
        chunk_index = 0
        # 表格不参与文本分块（避免表格文本混入到 text 块导致重复和顺序错乱）
        skip_categories = ['Header', 'Footer', 'PageBreak', 'Table']

        # section/父块标记
        current_section_id = 0
        current_parent_index = None

        def flush_current_chunk():
            nonlocal chunks, current_chunk_texts, current_length, chunk_index, current_chunk_start, current_chunk_end, current_section_id
            if not current_chunk_texts:
                return
            chunk_content = '\n'.join(current_chunk_texts)
            # 检查分块长度：如果超过 chunk_max，记录警告并截断
            # 注意：由于使用 '\n'.join()，实际长度可能比 current_length 大（因为添加了换行符）
            actual_length = len(chunk_content)
            if actual_length > chunk_max:
                logger.warning(
                    f"[CHUNK] 警告：分块长度 {actual_length} 超过限制 {chunk_max}（超出 {actual_length - chunk_max} 字符），"
                    f"chunk_index={chunk_index}, range=({current_chunk_start},{current_chunk_end})，"
                    f"文本片段数={len(current_chunk_texts)}，将截断到 {chunk_max} 字符"
                )
                chunk_content = chunk_content[:chunk_max]
                actual_length = chunk_max
            try:
                logger.debug(
                    f"[CHUNK][text] idx={chunk_index} sec={current_section_id} "
                    f"range=({current_chunk_start},{current_chunk_end}) len={len(chunk_content)}"
                )
            except Exception:
                pass
            chunks.append({
                'content': chunk_content,
                'chunk_index': chunk_index,
                'element_index_start': current_chunk_start,
                'element_index_end': current_chunk_end,
                'section_id': current_section_id,
                'is_parent': False,
            })
            chunk_index += 1
            current_chunk_texts = []
            current_length = 0
            current_chunk_start = None
            current_chunk_end = None

        def split_and_append(long_text: str, elem_idx: int):
            """将超长段落按 chunk_max 切分，逐段加入当前 section。
            每一小段的 element_index_start/end 都使用 elem_idx。
            严格控制：每个分块的长度不能超过 chunk_max。
            """
            nonlocal current_length, current_chunk_texts, current_chunk_start, current_chunk_end
            logger.debug(f"[SPLIT] elem_idx={elem_idx} long_len={len(long_text)} chunk_max={chunk_max}")
            i = 0
            while i < len(long_text):
                # 如果当前块已满，先刷新
                if current_length >= chunk_max:
                    flush_current_chunk()
                
                # 计算剩余空间（确保不超过 chunk_max）
                remain = chunk_max - current_length
                if remain <= 0:
                    # 这种情况理论上不应该发生（因为上面已经 flush），但为了安全还是处理
                    flush_current_chunk()
                    remain = chunk_max
                
                # 取剩余空间和剩余文本的最小值
                take = min(remain, len(long_text) - i)
                if take <= 0:
                    break
                
                piece = long_text[i:i + take]
                if not current_chunk_texts:
                    current_chunk_start = elem_idx
                current_chunk_texts.append(piece)
                current_chunk_end = elem_idx
                current_length += len(piece)
                
                # 再次检查：如果达到或超过上限，立即刷新（防止超过 chunk_max）
                if current_length >= chunk_max:
                    flush_current_chunk()
                
                i += take

        last_text_elem_idx: Optional[int] = None
        for element in elements:
            # 兼容 dict / obj 两种元素
            if isinstance(element, dict):
                category = element.get('category', 'NarrativeText')
                text = (element.get('text') or '').strip()
                elem_idx = element.get('element_index')
            else:
                category = getattr(element, 'category', 'NarrativeText')
                text = (getattr(element, 'text', '') or '').strip()
                elem_idx = getattr(element, 'element_index', None)
            if category in skip_categories:
                if category == 'Table':
                    logger.debug(f"[SKIP][table] elem_idx={elem_idx} text_len={len(text)}")
                continue
            if not text:
                continue
            text_length = len(text)
            if category in ['Title', 'ListItem']:
                # 如果当前块太小（< chunk_min），尝试继续累积而不是立即刷新
                # 这样可以减少小块数量，提高分块质量
                current_actual_length = len('\n'.join(current_chunk_texts)) if current_chunk_texts else 0
                if current_actual_length > 0 and current_actual_length < chunk_min:
                    # 当前块太小，尝试继续累积（如果添加后不超过 chunk_max）
                    newline_count = 1 if current_chunk_texts else 0
                    estimated_length = current_actual_length + text_length + newline_count
                    if estimated_length <= chunk_max:
                        # 可以继续累积，将 Title/ListItem 也加入当前块
                        logger.debug(f"[PARENT] 当前块太小({current_actual_length}<{chunk_min})，继续累积 Title/ListItem")
                        if not current_chunk_texts:
                            current_chunk_start = elem_idx
                        current_chunk_texts.append(text)
                        current_chunk_end = elem_idx
                        current_length = len('\n'.join(current_chunk_texts))
                        # 更新 section_id（但不创建新的父块）
                        current_section_id += 1
                        last_text_elem_idx = elem_idx
                        continue
                
                # 当前块足够大或无法继续累积，先刷新
                flush_current_chunk()
                # 新的父块（标题）
                current_section_id += 1
                logger.debug(f"[PARENT] idx={chunk_index} sec={current_section_id} elem_idx={elem_idx} len={len(text)}")
                chunks.append({
                    'content': text,
                    'chunk_index': chunk_index,
                    'element_index_start': elem_idx,
                    'element_index_end': elem_idx,
                    'section_id': current_section_id,
                    'is_parent': True,
                })
                current_parent_index = chunk_index
                chunk_index += 1
                # 重置累积器
                current_chunk_texts = []
                current_length = 0
                current_chunk_start = None
                current_chunk_end = None
            else:
                # 若与上一个文本元素不相邻（中间夹了表格或其他元素），先检查是否需要切断
                if last_text_elem_idx is not None and elem_idx is not None and current_chunk_texts:
                    if elem_idx != (current_chunk_end if current_chunk_end is not None else last_text_elem_idx) + 1:
                        # 元素索引不连续，但先检查当前块是否太小
                        current_actual_length = len('\n'.join(current_chunk_texts))
                        if current_actual_length < chunk_min:
                            # 当前块太小，尝试继续累积（即使索引不连续）
                            newline_count = 1 if current_chunk_texts else 0
                            estimated_length = current_actual_length + text_length + newline_count
                            if estimated_length <= chunk_max:
                                logger.debug(
                                    f"[BOUNDARY] 当前块太小({current_actual_length}<{chunk_min})，"
                                    f"继续累积（索引不连续：prev={last_text_elem_idx} current={elem_idx}）"
                                )
                                # 继续累积，不刷新
                            else:
                                # 无法继续累积，刷新
                                logger.debug(f"[BOUNDARY] flush by gap: prev={last_text_elem_idx} current={elem_idx}")
                                flush_current_chunk()
                        else:
                            # 当前块足够大，刷新
                            logger.debug(f"[BOUNDARY] flush by gap: prev={last_text_elem_idx} current={elem_idx}")
                            flush_current_chunk()

                if text_length > chunk_max:
                    split_and_append(text, elem_idx)
                else:
                    # 正常累积，必要时先换新块
                    # 计算添加后的实际长度：需要考虑换行符
                    # 如果 current_chunk_texts 不为空，添加新文本时会增加 1 个换行符
                    newline_count = 1 if current_chunk_texts else 0
                    estimated_length = current_length + text_length + newline_count
                    
                    # 严格检查：如果添加后超过 chunk_max，先刷新
                    if estimated_length > chunk_max:
                        flush_current_chunk()
                    
                    if not current_chunk_texts:
                        current_chunk_start = elem_idx
                    current_chunk_texts.append(text)
                    current_chunk_end = elem_idx
                    
                    # 重新计算实际长度（包括所有换行符），确保准确性
                    current_length = len('\n'.join(current_chunk_texts))
                    
                    # 再次检查：确保不会超过 chunk_max（理论上不应该，但为了安全）
                    if current_length > chunk_max:
                        logger.warning(
                            f"[CHUNK] 警告：分块长度 {current_length} 超过限制 {chunk_max}（超出 {current_length - chunk_max} 字符），"
                            f"element_index={elem_idx}, text_length={text_length}, "
                            f"文本片段数={len(current_chunk_texts)}，将刷新当前块"
                        )
                        # 如果超过限制，移除刚添加的文本并刷新
                        current_chunk_texts.pop()
                        flush_current_chunk()
                        # 重新添加（如果文本本身不超过 chunk_max）
                        if text_length <= chunk_max:
                            current_chunk_texts.append(text)
                            current_chunk_start = elem_idx
                            current_chunk_end = elem_idx
                            current_length = text_length
                        else:
                            # 如果文本本身超过 chunk_max，使用 split_and_append
                            split_and_append(text, elem_idx)
                last_text_elem_idx = elem_idx if elem_idx is not None else last_text_elem_idx
        # 收尾：仅在当前 section 内 flush，不做全文聚合
        flush_current_chunk()
        
        # 后处理：合并相邻的小块（如果合并后不超过 chunk_max）
        chunks = self._merge_small_chunks(chunks, chunk_max, chunk_min)
        
        logger.info(f"[DOCX] 结构分块完成，共 {len(chunks)} 个分块（section_id 最大为 {current_section_id}，chunk_max={chunk_max}，chunk_min={chunk_min}）")
        try:
            # 打印每个块的概览，便于定位错位/重复
            for c in chunks:
                logger.debug(
                    f"[CHUNK_SUMMARY] idx={c.get('chunk_index')} sec={c.get('section_id')} "
                    f"range=({c.get('element_index_start')},{c.get('element_index_end')}) "
                    f"len={len(c.get('content') or '')}"
                )
        except Exception:
            pass
        return chunks
    
    def _merge_small_chunks(self, chunks: List[Dict[str, Any]], chunk_max: int, chunk_min: int) -> List[Dict[str, Any]]:
        """合并相邻的小块（如果合并后不超过 chunk_max）"""
        if not chunks:
            return chunks
        
        merged_chunks: List[Dict[str, Any]] = []
        i = 0
        
        while i < len(chunks):
            current_chunk = chunks[i].copy()
            current_content = current_chunk.get('content', '')
            current_length = len(current_content)
            
            # 如果当前块足够大（>= chunk_min），直接添加
            if current_length >= chunk_min or current_chunk.get('is_parent', False):
                merged_chunks.append(current_chunk)
                i += 1
                continue
            
            # 当前块太小，尝试与后续块合并
            merged_content = current_content
            merged_start = current_chunk.get('element_index_start')
            merged_end = current_chunk.get('element_index_end')
            merged_section_id = current_chunk.get('section_id')
            j = i + 1
            
            # 尝试合并后续的小块
            while j < len(chunks):
                next_chunk = chunks[j]
                next_content = next_chunk.get('content', '')
                next_length = len(next_content)
                
                # 检查是否可以合并：
                # 1. 下一个块不是父块（Title/ListItem）
                # 2. 下一个块在同一个 section
                # 3. 合并后不超过 chunk_max
                if (not next_chunk.get('is_parent', False) and 
                    next_chunk.get('section_id') == merged_section_id):
                    # 计算合并后的长度（包括换行符）
                    merged_length = len(merged_content) + 1 + next_length  # +1 是换行符
                    
                    if merged_length <= chunk_max:
                        # 可以合并
                        merged_content = merged_content + '\n' + next_content
                        merged_end = next_chunk.get('element_index_end')
                        j += 1
                        continue
                
                # 无法继续合并，退出循环
                break
            
            # 创建合并后的块
            merged_chunk = {
                'content': merged_content,
                'chunk_index': len(merged_chunks),
                'element_index_start': merged_start,
                'element_index_end': merged_end,
                'section_id': merged_section_id,
                'is_parent': False,
            }
            merged_chunks.append(merged_chunk)
            
            # 跳过已合并的块
            i = j
        
        # 更新 chunk_index
        for idx, chunk in enumerate(merged_chunks):
            chunk['chunk_index'] = idx
        
        if len(merged_chunks) < len(chunks):
            logger.info(
                f"[CHUNK_MERGE] 合并小块完成：原始分块数={len(chunks)}，"
                f"合并后分块数={len(merged_chunks)}，减少了 {len(chunks) - len(merged_chunks)} 个分块"
            )
        
        return merged_chunks
